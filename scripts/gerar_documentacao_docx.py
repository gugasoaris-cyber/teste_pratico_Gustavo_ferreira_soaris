import sys
from pathlib import Path

import markdown
from bs4 import BeautifulSoup, NavigableString
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor

RAIZ = Path(__file__).resolve().parents[1]
MD = RAIZ / "DOCUMENTACAO.md"
SAIDA = RAIZ / "DOCUMENTACAO.docx"


def _adicionar_runs(paragraph, node):
    for child in node.children:
        if isinstance(child, NavigableString):
            t = str(child)
            if t:
                paragraph.add_run(t)
        elif child.name in ("strong", "b"):
            r = paragraph.add_run(child.get_text())
            r.bold = True
        elif child.name in ("em", "i"):
            r = paragraph.add_run(child.get_text())
            r.italic = True
        elif child.name == "code":
            r = paragraph.add_run(child.get_text())
            r.font.name = "Consolas"
            r.font.size = Pt(9)
        elif child.name == "br":
            paragraph.add_run().add_break()
        elif child.name == "a":
            paragraph.add_run(child.get_text())
        else:
            _adicionar_runs(paragraph, child)


def _paragrafo_de_elemento(doc, elem):
    p = doc.add_paragraph()
    _adicionar_runs(p, elem)


def _definir_borda_tabela(tbl):
    tbl_pr = tbl._tbl.tblPr
    if tbl_pr is None:
        from docx.oxml import parse_xml

        tbl_pr = parse_xml(
            r'<w:tblPr xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"/>'
        )
        tbl._tbl.insert(0, tbl_pr)
    borders = OxmlElement("w:tblBorders")
    for nome in ("top", "left", "bottom", "right", "insideH", "insideV"):
        b = OxmlElement(f"w:{nome}")
        b.set(qn("w:val"), "single")
        b.set(qn("w:sz"), "4")
        b.set(qn("w:space"), "0")
        b.set(qn("w:color"), "CCCCCC")
        borders.append(b)
    tbl_pr.append(borders)


def _processar_elemento(doc, elem):
    if not getattr(elem, "name", None):
        return
    nome = elem.name.lower()
    if nome == "h1":
        doc.add_heading(elem.get_text().strip(), level=0)
    elif nome == "h2":
        doc.add_heading(elem.get_text().strip(), level=1)
    elif nome == "h3":
        doc.add_heading(elem.get_text().strip(), level=2)
    elif nome == "h4":
        doc.add_heading(elem.get_text().strip(), level=3)
    elif nome == "p":
        _paragrafo_de_elemento(doc, elem)
    elif nome == "hr":
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after = Pt(6)
        p.add_run("_" * 60).font.color.rgb = RGBColor(180, 180, 180)
    elif nome == "pre":
        texto = elem.get_text()
        p = doc.add_paragraph()
        r = p.add_run(texto)
        r.font.name = "Consolas"
        r.font.size = Pt(9)
        p.paragraph_format.left_indent = Pt(18)
    elif nome == "ul":
        for li in elem.find_all("li", recursive=False):
            p = doc.add_paragraph(style="List Bullet")
            _adicionar_runs(p, li)
    elif nome == "ol":
        for li in elem.find_all("li", recursive=False):
            p = doc.add_paragraph(style="List Number")
            _adicionar_runs(p, li)
    elif nome == "table":
        linhas = elem.find_all("tr")
        if not linhas:
            return
        colunas = len(linhas[0].find_all(["th", "td"]))
        tbl = doc.add_table(rows=len(linhas), cols=max(colunas, 1))
        tbl.style = "Table Grid"
        _definir_borda_tabela(tbl)
        for i, tr in enumerate(linhas):
            celulas = tr.find_all(["th", "td"])
            for j, td in enumerate(celulas):
                if j >= len(tbl.rows[i].cells):
                    break
                cell = tbl.rows[i].cells[j]
                cell.text = ""
                p = cell.paragraphs[0]
                for sub in td.children:
                    if isinstance(sub, NavigableString):
                        p.add_run(str(sub))
                    elif getattr(sub, "name", None):
                        _adicionar_runs(p, sub)
                if td.name == "th":
                    for r in p.runs:
                        r.bold = True
    else:
        if elem.get_text(strip=True):
            _paragrafo_de_elemento(doc, elem)


def html_para_docx(html: str, caminho_saida: Path):
    soup = BeautifulSoup(html, "html.parser")
    doc = Document()
    estilo_normal = doc.styles["Normal"]
    estilo_normal.font.name = "Calibri"
    estilo_normal.font.size = Pt(11)

    body = soup.body
    if body is None:
        body = soup
    for elem in body.children:
        if getattr(elem, "name", None):
            _processar_elemento(doc, elem)

    doc.save(caminho_saida)


def principal():
    if not MD.exists():
        print(f"Ficheiro não encontrado: {MD}", file=sys.stderr)
        sys.exit(1)
    texto = MD.read_text(encoding="utf-8")
    html = markdown.markdown(
        texto,
        extensions=["tables", "fenced_code", "nl2br"],
    )
    html_completo = f"<html><body>{html}</body></html>"
    html_para_docx(html_completo, SAIDA)
    print(f"Gerado: {SAIDA}")


if __name__ == "__main__":
    principal()
